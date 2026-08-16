from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "Impacto_y_Estrategia_Comercial_Lucia_Movistar.docx"
LOGO = ROOT / "frontend" / "logo.png"

BLUE = "00A9E0"
BLUE_DARK = "0081C9"
NAVY = "0B2739"
INK = "183747"
MUTED = "617784"
PALE_BLUE = "EAF8FD"
PALE_GREEN = "EAF8F4"
PALE_GRAY = "F4F7F9"
WHITE = "FFFFFF"
GREEN = "07956D"
GOLD = "D98E04"
RED = "B23A48"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_font(run, size=None, bold=None, color=INK, name="Calibri", italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = rgb(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=140):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def border_paragraph(paragraph, color=BLUE, size="18"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def shade_paragraph(paragraph, fill, border=None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if border:
        p_bdr = OxmlElement("w:pBdr")
        for side in ("top", "left", "bottom", "right"):
            edge = OxmlElement(f"w:{side}")
            edge.set(qn("w:val"), "single")
            edge.set(qn("w:sz"), "8")
            edge.set(qn("w:color"), border)
            edge.set(qn("w:space"), "6")
            p_bdr.append(edge)
        p_pr.append(p_bdr)


def add_numbering_definition(doc, ordered=False):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abs + [0]) + 1
    num_id = max(existing_num + [0]) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "%1." if ordered else "•")
    level.extend([start, fmt, text])
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "280")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "290")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_item(doc, text, num_id, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    set_font(p.add_run("• "), color=BLUE_DARK)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_font(rest)
    else:
        set_font(p.add_run(text))
    return p


def add_body(doc, text, bold_lead=None, after=8, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    p.alignment = align or WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_font(r2)
    else:
        set_font(p.add_run(text))
    return p


def add_callout(doc, label, text, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    set_font(p.add_run(f"{label.upper()}  "), size=9, bold=True, color=accent)
    set_font(p.add_run(text), size=11, bold=True, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(text), size={1: 16, 2: 13, 3: 12}[level], bold=True,
             color={1: BLUE_DARK, 2: BLUE_DARK, 3: NAVY}[level])
    return p


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper())
    set_font(r, size=9, bold=True, color=BLUE)
    return p


def add_table(doc, headers, rows, widths_dxa, header_fill=NAVY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths_dxa)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, value in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(value), size=9.5, bold=True, color=WHITE)
    for row_index, row_values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cell = cells[idx]
            set_cell_shading(cell, WHITE if row_index % 2 == 0 else PALE_GRAY)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            set_font(p.add_run(str(value)), size=9.4, color=INK)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.78)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

# Narrative proposal token map with named branded overrides:
# compact 0.78/0.72 vertical margins, Movistar blue palette, proposal centerpiece.
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(10.5)
normal.font.color.rgb = rgb(INK)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE_DARK, 18, 10),
    ("Heading 2", 13, BLUE_DARK, 12, 6),
    ("Heading 3", 12, NAVY, 8, 4),
):
    style = doc.styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = rgb(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(hp.add_run("LUCÍA  |  ASISTENTE INTELIGENTE DE FACTURACIÓN"), size=8.5, bold=True, color=MUTED)
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_font(fp.add_run("Propuesta de impacto y estrategia comercial  •  Agosto 2026"), size=8.5, color=MUTED)
add_page_number(footer.add_paragraph())

bullet_id = add_numbering_definition(doc, ordered=False)
number_id = add_numbering_definition(doc, ordered=True)

# Cover
if LOGO.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(36)
    run = p.add_run()
    run.add_picture(str(LOGO), width=Inches(0.78))
    inline = run._element.xpath('.//wp:docPr')
    if inline:
        inline[0].set('descr', 'Logotipo de Movistar')

for _ in range(2):
    doc.add_paragraph()
add_kicker(doc, "Propuesta de impacto y estrategia comercial")
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(8)
set_font(title.add_run("Lucía"), size=34, bold=True, color=NAVY)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(20)
set_font(subtitle.add_run("La inteligencia que convierte un recibo complejo\nen una explicación que el cliente sí entiende"), size=17, bold=True, color=BLUE_DARK)
lead = doc.add_paragraph()
lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
lead.paragraph_format.space_after = Pt(30)
set_font(lead.add_run("Asistente híbrido de facturación, privacidad por sesión, trazabilidad y experiencia visual personalizada."), size=11.5, color=MUTED)
border_paragraph(lead, BLUE, "22")
add_callout(doc, "Promesa central", "Menos confusión, respuestas más confiables y una ruta clara desde la duda hasta la resolución.")
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(38)
set_font(meta.add_run("Documento comercial y estratégico  |  MVP funcional  |  Perú"), size=10, bold=True, color=MUTED)

doc.add_page_break()

# 1
add_kicker(doc, "01  |  Visión ejecutiva")
add_heading(doc, "Un producto para recuperar la confianza en la facturación", 1)
add_body(doc, "Lucía transforma la consulta de un recibo en una experiencia guiada: reconoce la intención del cliente, recupera únicamente la información autorizada, calcula los cambios con reglas verificables y utiliza inteligencia artificial para adaptar la explicación al lenguaje de cada persona.")
add_callout(doc, "Tesis de negocio", "La oportunidad no consiste en responder más mensajes, sino en resolver mejor las dudas que generan reclamos, llamadas repetidas y desconfianza.", PALE_BLUE, BLUE_DARK)

add_heading(doc, "El problema que resuelve", 2)
for text, lead_text in [
    ("Recibos difíciles de interpretar. Cargos, descuentos, bonificaciones, prorrateos y nombres internos obligan al cliente a descifrar su propia factura.", "Recibos difíciles de interpretar."),
    ("Atención costosa y repetitiva. Una pregunta simple puede terminar en varios contactos, derivaciones o una llamada al asesor.", "Atención costosa y repetitiva."),
    ("Chatbots que pierden contexto. Las respuestas rígidas repiten bloques, ignoran seguimientos breves y dañan la percepción de inteligencia.", "Chatbots que pierden contexto."),
    ("Riesgo de respuestas no sustentadas. Un modelo generativo sin controles puede confundir total facturado con deuda o inventar causas.", "Riesgo de respuestas no sustentadas."),
    ("Información sensible expuesta. Una experiencia sin identidad y autorización adecuadas puede cruzar datos entre clientes.", "Información sensible expuesta."),
]:
    add_list_item(doc, text, bullet_id, lead_text)

add_heading(doc, "La solución en una frase", 2)
add_body(doc, "Lucía es una capa de comprensión y confianza sobre los datos de facturación: explica, compara, visualiza y deriva cuando corresponde, sin sustituir la fuente oficial ni afirmar lo que los datos no permiten comprobar.")

doc.add_page_break()

# 2
add_kicker(doc, "02  |  Impacto")
add_heading(doc, "Impacto para el cliente, la operación y la marca", 1)
add_heading(doc, "1. Experiencia del cliente", 2)
for text in [
    "Traduce conceptos técnicos a lenguaje cotidiano y adapta la extensión: una línea, explicación simple o desglose completo.",
    "Mantiene el contexto ante seguimientos como “¿por qué?”, “más fácil” o “muéstrame todas”.",
    "Presenta total, vencimiento, variación, causa principal e historial en un panel personal e interactivo.",
    "Agrupa bonos y cargos que se compensan para mostrar su efecto real, evitando sumas engañosas.",
]:
    add_list_item(doc, text, bullet_id)

add_heading(doc, "2. Eficiencia operativa", 2)
for text in [
    "Resuelve de forma autónoma preguntas frecuentes sobre montos, vencimientos, cambios, bonos, reconexiones e historial.",
    "Reduce el tiempo que un asesor dedica a reconstruir el caso porque la derivación conserva el contexto útil.",
    "Separa consultas generales de consultas personales: el cliente puede conversar sin autenticarse y solo inicia sesión cuando necesita datos privados.",
    "Registra métricas de interacción, derivación, satisfacción y duración para aprender del uso real.",
]:
    add_list_item(doc, text, bullet_id)

add_heading(doc, "3. Confianza y reputación", 2)
for text in [
    "Distingue total facturado, estado de deuda y saldo pendiente; no presenta esos conceptos como equivalentes.",
    "Reconoce la ausencia o contradicción de un dato y lo comunica en lugar de completar el vacío con una suposición.",
    "Limita cada consulta a la identidad autenticada y elimina el contexto privado al cerrar sesión.",
]:
    add_list_item(doc, text, bullet_id)

add_callout(doc, "Impacto esperado en piloto", "Mayor comprensión del recibo, más resolución digital, menos derivaciones evitables y mejor satisfacción. Estas variables deben medirse contra una línea base; todavía no se presentan como resultados productivos.", PALE_GREEN, GREEN)

doc.add_page_break()

# 3
add_kicker(doc, "03  |  Marco diferenciador")
add_heading(doc, "Por qué Lucía no es “otro chatbot”", 1)
add_table(doc,
          ["Dimensión", "Chatbot convencional", "Lucía"],
          [
              ("Forma de responder", "Flujos rígidos o texto generativo genérico.", "Respuesta adaptativa con contexto, tono y nivel de detalle."),
              ("Cálculo financiero", "Puede quedar delegado al modelo.", "Motor determinístico: la IA no calcula ni inventa importes."),
              ("Fuente de verdad", "Documentos amplios sin trazabilidad suficiente.", "Datos del cliente, factura, servicio y eventos asociados."),
              ("Privacidad", "Validación tardía o solo en la interfaz.", "Autorización en servidor y aislamiento por sesión autenticada."),
              ("Experiencia", "Burbuja de texto repetitiva.", "Dashboard personal, recibo expandido, tarjetas y evolución visual."),
              ("Escalamiento", "Deriva sin contexto.", "Handoff con motivo y conversación útil para el asesor."),
              ("Medición", "Volumen de chats.", "Resolución proxy, derivación, satisfacción, repetición y duración."),
          ], [1800, 3300, 4260])

add_heading(doc, "El modelo híbrido", 2)
add_body(doc, "La ventaja técnica y comercial está en separar responsabilidades. La capa determinística encuentra el cliente correcto, selecciona la factura, suma cargos, compara periodos y valida eventos. La capa de IA interpreta lenguaje natural, jerga, errores de escritura y preferencias de formato. Finalmente, una barrera de seguridad evita que la redacción introduzca cifras, fechas o facturas que no estaban en el borrador verificado.")
add_callout(doc, "Fórmula diferenciadora", "Datos autorizados + cálculo verificable + IA adaptativa + experiencia visual + derivación con contexto.")

add_callout(doc, "Activos defendibles", "Conocimiento especializado en facturación, biblioteca de escenarios y pruebas de regresión, y una arquitectura reutilizable en otras industrias con estados de cuenta complejos.", PALE_GREEN, GREEN)

doc.add_page_break()

# 4
add_kicker(doc, "04  |  Producto")
add_heading(doc, "Cómo funciona la experiencia", 1)
steps = [
    ("El cliente conversa", "Puede hacer preguntas generales sin iniciar sesión y utiliza su forma natural de hablar."),
    ("La intención se interpreta", "El sistema identifica si busca un monto, vencimiento, causa, historial, plan o asesor."),
    ("Los datos se verifican", "Las reglas consultan únicamente la cuenta autorizada y calculan importes desde los registros."),
    ("La respuesta se adapta", "La IA simplifica, resume o amplía sin modificar los hechos financieros permitidos."),
    ("La interfaz explica", "Tarjetas, comparación, cargos agrupados y gráfico histórico convierten datos en comprensión."),
    ("El caso escala si es necesario", "El asesor recibe el motivo y el contexto, evitando que el cliente repita todo."),
]
for title_text, detail in steps:
    add_list_item(doc, f"{title_text}. {detail}", number_id, f"{title_text}.")

add_heading(doc, "Capacidades demostrables del MVP", 2)
for text in [
    "Registro e inicio de sesión con IDs existentes en los datos; sesiones persistentes y cierre seguro.",
    "Privacidad por cliente: rechazo de consultas sobre otra cuenta, incluso si se manipulan parámetros.",
    "Consulta del último recibo, vencimiento, estado, historial y comparación entre periodos.",
    "Explicación de aumentos, bonos, descuentos, reconexiones y prorrateos con evidencia.",
    "Catálogo de planes con memoria conversacional para seguimientos como “todas”.",
    "Dashboard personal, detalle expandido, cargos compensados y gráfico animado.",
    "Derivación a asesor, encuesta de satisfacción y dashboard de métricas.",
    "Suite técnica actual: 87 pruebas automatizadas aprobadas en el entorno local del MVP.",
]:
    add_list_item(doc, text, bullet_id)

add_callout(doc, "Límite honesto", "El MVP no debe venderse todavía como una plataforma productiva completa: requiere integración con identidad corporativa, pagos en tiempo real, observabilidad, seguridad empresarial y evaluación con datos etiquetados.", "FFF5E6", GOLD)

doc.add_page_break()

# 5
add_kicker(doc, "05  |  Mercado y posicionamiento")
add_heading(doc, "Dónde encaja y quién lo compra", 1)
add_body(doc, "El punto de entrada natural son organizaciones con facturas o estados de cuenta difíciles de explicar y un volumen alto de contactos repetitivos. La propuesta se puede posicionar como una solución B2B2C de explicación inteligente, integrada en el canal digital existente.")
add_table(doc,
          ["Comprador", "Dolor principal", "Valor que debe escuchar"],
          [
              ("Experiencia de cliente", "Baja comprensión y satisfacción.", "Explicaciones personalizadas que elevan la claridad y confianza."),
              ("Operaciones / Contact center", "Contactos repetidos y alto tiempo de atención.", "Automatización segura y derivaciones con contexto."),
              ("Facturación", "Reclamos por variaciones y conceptos técnicos.", "Causas respaldadas y conciliación visible de cargos."),
              ("Canales digitales", "Bajo uso efectivo del autoservicio.", "Una experiencia integrada, visual y medible."),
              ("Riesgo / Seguridad", "Exposición de datos o respuestas no sustentadas.", "Aislamiento por identidad, trazabilidad y límites explícitos."),
              ("Tecnología / Datos", "Integraciones fragmentadas.", "Capa modular conectable a APIs, data lake y sistemas existentes."),
          ], [2100, 3300, 3960])

add_heading(doc, "Posicionamiento recomendado", 2)
add_callout(doc, "Categoría", "Plataforma de explicación inteligente y verificable para facturas y estados de cuenta.")
add_body(doc, "No conviene venderla como “un chatbot con IA”. Esa descripción la vuelve comparable con herramientas genéricas. Debe presentarse como una solución de resolución digital especializada, capaz de convertir datos transaccionales complejos en respuestas comprensibles y seguras.")

add_heading(doc, "Mercados adyacentes", 2)
add_body(doc, "Después de telecomunicaciones, el mismo patrón se puede aplicar a energía y agua, tarjetas y préstamos, seguros, educación, salud privada y suscripciones. En todos ellos existe una combinación de movimientos, descuentos, periodos, saldo, vencimientos y preguntas repetitivas.")

doc.add_page_break()

# 6
add_kicker(doc, "06  |  Oferta y monetización")
add_heading(doc, "Cómo empaquetar y cobrar el producto", 1)
add_table(doc,
          ["Componente", "Qué incluye", "Cómo monetizar"],
          [
              ("Descubrimiento", "Mapa de datos, casos, riesgos y línea base.", "Tarifa fija por diagnóstico."),
              ("Piloto controlado", "Integración limitada, escenarios prioritarios y medición.", "Tarifa fija por implementación del piloto."),
              ("Plataforma", "Motor híbrido, sesiones, seguridad, panel y métricas.", "Suscripción anual por organización o marca."),
              ("Uso", "Conversaciones o explicaciones procesadas.", "Cargo variable por interacción o tramo de volumen."),
              ("Módulos", "Handoff, analítica avanzada, multicanal, idiomas, NBO.", "Licencia adicional por módulo."),
              ("Servicios", "Integración, afinamiento, capacitación y soporte.", "Bolsa de horas o servicio administrado."),
          ], [1800, 4300, 3260])

add_heading(doc, "Tres paquetes comerciales", 2)
for title_text, detail in [
    ("Piloto de valor", "Uno o dos productos, conjunto acotado de causas, usuarios controlados y tablero de evaluación."),
    ("Operación empresarial", "Integraciones productivas, autenticación corporativa, monitoreo, SLA y gobierno de cambios."),
    ("Plataforma expandida", "Múltiples unidades de negocio, canales, idiomas y modelos comerciales sobre la misma capa de explicación."),
]:
    add_list_item(doc, f"{title_text}. {detail}", bullet_id, f"{title_text}.")

add_heading(doc, "Lógica de precio", 2)
add_body(doc, "El precio debe anclarse al valor económico, no al costo de tokens. La conversación comercial debe estimar contactos evitados, minutos ahorrados, costo promedio por atención, reducción de repetición y mejora de adopción digital. La tarifa se diseña para capturar una fracción del beneficio esperado, dejando un retorno visible para el comprador.")
add_callout(doc, "Fórmula de ROI", "Beneficio mensual = contactos evitados × costo por contacto + minutos ahorrados × costo por minuto + valor de retención atribuible. ROI = (beneficio − costo total) / costo total.", PALE_GREEN, GREEN)

doc.add_page_break()

# 7
add_kicker(doc, "07  |  Estrategia de venta")
add_heading(doc, "Cómo vender Lucía", 1)
add_heading(doc, "Movimiento comercial recomendado", 2)
sales_steps = [
    ("Entrar por el dolor", "Seleccionar una causa de alto volumen: aumento del recibo, deuda, prorrateo o reconexión."),
    ("Demostrar con un caso real", "Mostrar la factura técnica, la respuesta habitual y la explicación de Lucía lado a lado."),
    ("Vender un piloto medible", "Acordar población, línea base, métricas, duración y criterios de éxito antes de integrar."),
    ("Probar seguridad y precisión", "Ejecutar pruebas de aislamiento, montos, historial y ausencia de afirmaciones sin evidencia."),
    ("Escalar por módulos", "Añadir servicios, canales, idiomas, handoff y recomendaciones cuando el núcleo ya demuestre valor."),
]
for title_text, detail in sales_steps:
    add_list_item(doc, f"{title_text}. {detail}", number_id, f"{title_text}.")

add_heading(doc, "Discurso de 30 segundos", 2)
quote = doc.add_paragraph()
quote.paragraph_format.left_indent = Inches(0.25)
quote.paragraph_format.right_indent = Inches(0.25)
quote.paragraph_format.space_before = Pt(6)
quote.paragraph_format.space_after = Pt(14)
quote.paragraph_format.line_spacing = 1.25
shade_paragraph(quote, PALE_BLUE, BLUE)
set_font(quote.add_run("“Lucía ayuda a tus clientes a entender exactamente qué cambió en su recibo y por qué. Combina reglas financieras verificables con IA conversacional, protege los datos por identidad y entrega una experiencia visual que puede resolver dudas antes de que se conviertan en llamadas o reclamos.”"), size=12.5, italic=True, color=NAVY)

add_heading(doc, "Demostración comercial ideal", 2)
for text in [
    "Abrir una conversación sin sesión y hacer una consulta general.",
    "Iniciar sesión y mostrar el panel privado del recibo.",
    "Preguntar “¿por qué subió?” y continuar con “más fácil”, “¿desde cuándo?” y “¿por qué?”.",
    "Abrir el recibo completo y enseñar cargos compensados e historial.",
    "Intentar consultar otro cliente para demostrar el bloqueo de privacidad.",
    "Cerrar sesión y confirmar que desaparecen mensajes y datos privados.",
]:
    add_list_item(doc, text, number_id)

doc.add_page_break()

# 8
add_kicker(doc, "08  |  Piloto, medición y cierre")
add_heading(doc, "Plan de piloto propuesto", 1)
add_table(doc,
          ["Fase", "Objetivo", "Entregable", "Criterio de avance"],
          [
              ("1. Descubrimiento", "Priorizar causas y fuentes.", "Mapa de datos y línea base.", "Datos y dueños confirmados."),
              ("2. Integración", "Conectar identidad y facturación.", "Flujos seguros de prueba.", "Casos trazables extremo a extremo."),
              ("3. Afinamiento", "Optimizar respuestas y UX.", "Biblioteca de escenarios.", "Precisión acordada en muestra etiquetada."),
              ("4. Piloto", "Medir comportamiento real.", "Tablero y análisis de resultados.", "Cumplimiento de umbrales definidos."),
              ("5. Escala", "Industrializar la solución.", "Arquitectura, SLA y roadmap.", "Caso de negocio aprobado."),
          ], [1450, 2500, 2700, 2710])

add_heading(doc, "Indicadores que deben gobernar la decisión", 2)
add_table(doc,
          ["Indicador", "Qué demuestra", "Cómo medirlo"],
          [
              ("Comprensión", "El cliente entiende monto y causa.", "Pregunta posterior o prueba breve de comprensión."),
              ("Resolución digital", "La duda termina sin asesor.", "Interacciones resueltas / interacciones elegibles."),
              ("Recontacto", "La explicación evita repetir la consulta.", "Nuevo contacto por misma causa y cliente."),
              ("Derivación correcta", "El sistema escala cuando corresponde.", "Casos etiquetados y revisión del asesor."),
              ("Exactitud financiera", "No cambia montos ni causas.", "Comparación contra ground truth de facturación."),
              ("Satisfacción", "La experiencia genera confianza.", "Encuesta posterior y tasa de respuesta."),
              ("Tiempo de atención", "Reduce esfuerzo operativo.", "Duración digital y minutos evitados al asesor."),
          ], [2200, 3300, 3860])

doc.add_page_break()
add_kicker(doc, "09  |  Preparación para producción")
add_heading(doc, "Riesgos que deben resolverse antes de producción", 1)
for text in [
    "Integración con autenticación e identidad corporativa, sin credenciales locales de demostración.",
    "Acceso a saldo pendiente y pagos actualizados para separar factura, deuda y aplicación del pago.",
    "Observabilidad, retención de logs, protección de secretos, gestión de consentimientos y pruebas de penetración.",
    "Evaluación continua con casos etiquetados para medir recuperación, alucinación financiera y derivación.",
    "Gobierno de catálogo, cambios de esquema y responsables de calidad de datos.",
]:
    add_list_item(doc, text, bullet_id)

add_callout(doc, "Recomendación final", "Vender primero un piloto enfocado en explicación de aumentos y composición del recibo. Es el caso donde Lucía muestra con mayor claridad su diferenciación: entiende el lenguaje, verifica la matemática, protege la identidad y convierte datos complejos en una experiencia visual.")

closing = doc.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
closing.paragraph_format.space_before = Pt(18)
closing.paragraph_format.space_after = Pt(4)
set_font(closing.add_run("Lucía no vende respuestas. Vende comprensión, confianza y resolución."), size=15, bold=True, color=NAVY)

# Metadata and document hygiene
doc.core_properties.title = "Impacto y estrategia comercial de Lucía"
doc.core_properties.subject = "Propuesta de valor, diferenciación y comercialización"
doc.core_properties.author = "Equipo del proyecto Lucía"
doc.core_properties.keywords = "facturación, inteligencia artificial, experiencia de cliente, telecomunicaciones, estrategia comercial"
doc.core_properties.comments = "Documento comercial del MVP"

doc.save(OUTPUT)
print(OUTPUT)
